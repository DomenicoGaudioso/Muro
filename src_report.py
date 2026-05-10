# -*- coding: utf-8 -*-
"""
src_report.py — Generazione Relazione Tecnica Muro di Sostegno
Formati: Word (.docx) e PDF (.pdf)
NTC2018 §6.5.3 / EC7-1 §6.5 / EC2 §6
"""
from __future__ import annotations

import io
import math
import tempfile
from datetime import date
from pathlib import Path
from typing import Dict, Any

# ---------------------------------------------------------------------------
# WORD REPORT
# ---------------------------------------------------------------------------

def create_word_report(data: Dict[str, Any]) -> bytes:
    """
    Genera la relazione tecnica in formato Word (.docx).
    Restituisce i bytes del file, senza scrivere su disco.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _aggiungi_heading1(doc, testo):
        p = doc.add_paragraph()
        run = p.add_run(testo)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(8)
        return p

    def _aggiungi_heading2(doc, testo):
        p = doc.add_paragraph()
        run = p.add_run(testo)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        return p

    def _intestazione_tabella(tabella):
        for cell in tabella.rows[0].cells:
            for par in cell.paragraphs:
                for run in par.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'), 'E5E7EB')
            cell._tc.get_or_add_tcPr().append(shading)

    def _esito_cell(cell, fs_value, limite=1.0):
        par = cell.paragraphs[0]
        run = par.add_run()
        if fs_value >= limite:
            run.text = f"OK {fs_value:.3f}"
        else:
            run.text = f"NO {fs_value:.3f}"
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        run.bold = True

    doc = Document()
    r = data['risultati']
    warnings = data.get("warnings", [])
    q_ratio_statico = r['statico']['qmax'] / max(data['q_amm'], 1e-9)

    def _esito_testo(valore, limite=1.0, inverse=False):
        ok = valore <= limite if inverse else valore >= limite
        return "VERIFICATO" if ok else "NON VERIFICATO"

    def _aggiungi_tabella_kv(doc, righe):
        tabella = doc.add_table(rows=1, cols=2)
        tabella.style = 'Table Grid'
        tabella.rows[0].cells[0].text = "Voce"
        tabella.rows[0].cells[1].text = "Valore"
        _intestazione_tabella(tabella)
        for label, value in righe:
            row = tabella.add_row().cells
            row[0].text = str(label)
            row[1].text = str(value)
        return tabella

    def _dati_muro():
        from src import DatiMuro, DEFAULT_STRAT

        return DatiMuro(
            H=data["H"],
            q=data["q"],
            gamma_cls=data["gamma_cls"],
            B=data["B"],
            B_punta=data["B_punta"],
            B_tallone=data["B_tallone"],
            t_base=data["t_base"],
            t_fusto_top=data["t_fusto_top"],
            t_fusto_bot=data["t_fusto_bot"],
            mu_base=data["mu_base"],
            q_amm=data["q_amm"],
            h_fronte=data["h_fronte"],
            falda_retro=data["falda_retro"],
            falda_fronte=data["falda_fronte"],
            kh=data["kh"],
            kv=data["kv"],
            delta_muro=data["delta_muro"],
            include_passivo=data.get("include_passivo", True),
            stratigrafia_csv=data.get("stratigrafia_csv", DEFAULT_STRAT),
            ha_tirante=data.get("ha_tirante", False),
            t_quota=data.get("t_quota", 0.0),
            t_inclinazione=data.get("t_inclinazione", 0.0),
            t_tiro=data.get("t_tiro", 0.0),
            analizza_pendio=bool(data.get("pendio")),
            pendio_H=data.get("pendio_H", 5.0),
            pendio_beta=data.get("pendio_beta", 26.0),
            pendio_berma=data.get("pendio_berma", 6.0),
        )

    def _plot_png_bytes(fig, width=1100, height=700):
        import plotly.io as pio

        fig.update_layout(font=dict(color="#000000"), paper_bgcolor="white", plot_bgcolor="white")
        return pio.to_image(fig, format="png", width=width, height=height, scale=2)

    def _aggiungi_immagine(doc, titolo, png_bytes, width_cm=16.0):
        _aggiungi_heading2(doc, titolo)
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            tmp.write(png_bytes)
            tmp_path = tmp.name
        try:
            doc.add_picture(tmp_path, width=Cm(width_cm))
            last = doc.paragraphs[-1]
            last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    # Margini
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    _aggiungi_heading1(doc, "1. RELAZIONE DI CALCOLO - MURO DI SOSTEGNO")

    _aggiungi_heading2(doc, "1.1 Frontespizio")
    p = doc.add_paragraph()
    p.add_run("Applicazione: ").bold = True
    p.add_run("Muro - Analisi Avanzata NTC2018")
    p = doc.add_paragraph()
    p.add_run("Normative di riferimento: ").bold = True
    p.add_run("NTC2018 §6.5.3, EC7-1 §6.5, EC8-5 §7.3")
    p = doc.add_paragraph()
    p.add_run("Data elaborazione: ").bold = True
    p.add_run(str(date.today()))
    p = doc.add_paragraph()
    p.add_run("Disclaimer: ").bold = True
    p.add_run(
        "I risultati prodotti dall'applicazione devono essere verificati da un ingegnere "
        "abilitato prima di essere utilizzati per qualsiasi scopo progettuale o costruttivo."
    )
    doc.add_paragraph()

    _aggiungi_heading2(doc, "1.2 Sintesi esecutiva")
    _aggiungi_tabella_kv(
        doc,
        [
            ("Ribaltamento statico EQU", f"{_esito_testo(r['st_EQU']['FS_rib'])} - FS = {r['st_EQU']['FS_rib']:.3f}"),
            ("Scorrimento statico GEO", f"{_esito_testo(r['statico']['FS_scorr'])} - FS = {r['statico']['FS_scorr']:.3f}"),
            ("Portanza statica GEO", f"{_esito_testo(r['statico']['FS_portanza'])} - FS = {r['statico']['FS_portanza']:.3f}"),
            ("Pressione massima / q_amm", f"{_esito_testo(q_ratio_statico, 1.0, inverse=True)} - {q_ratio_statico:.3f}"),
            ("Criticita automatiche", f"{len(warnings)}"),
        ],
    )
    if warnings:
        _aggiungi_heading2(doc, "Criticita da esaminare")
        for warning in warnings:
            doc.add_paragraph(str(warning), style="List Bullet")

    _aggiungi_heading2(doc, "Riferimenti e perimetro del modello")
    for ref in [
        "NTC 2018, Capitolo 6.5 - opere di sostegno: azioni, drenaggio, riempimenti e verifiche del complesso opera-terreno.",
        "Circolare 21/01/2019 n. 7 C.S.LL.PP., C6.5.3.1.1: SLU GEO per ribaltamento, scorrimento, carico limite e stabilita globale.",
        "EN 1997-1, Sezione 9: stati limite delle strutture di sostegno, pressioni idrauliche, scorrimento, ribaltamento e portanza.",
    ]:
        doc.add_paragraph(ref, style="List Bullet")
    doc.add_paragraph(
        "La relazione documenta le verifiche automatiche principali; non sostituisce la relazione geologica, "
        "la scelta dei parametri caratteristici ne la verifica indipendente del progettista."
    )

    _aggiungi_heading2(doc, "1.3 Dati di input")

    _aggiungi_heading2(doc, "1.3.1 Geometria")
    t = doc.add_table(rows=1, cols=4)
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    hdr[0].text = "Parametro"
    hdr[1].text = "Valore"
    hdr[2].text = "Unità"
    hdr[3].text = "Descrizione"
    _intestazione_tabella(t)

    geo_rows = [
        ("H", f"{data['H']:.2f}", "m", "Altezza terreno a tergo"),
        ("B", f"{data['B']:.2f}", "m", "Base totale fondazione"),
        ("B_punta", f"{data['B_punta']:.2f}", "m", "Larghezza punta"),
        ("B_tallone", f"{data['B_tallone']:.2f}", "m", "Larghezza tallone"),
        ("t_base", f"{data['t_base']:.2f}", "m", "Spessore base"),
        ("t_fusto_top", f"{data['t_fusto_top']:.2f}", "m", "Spessore fusto in testa"),
        ("t_fusto_bot", f"{data['t_fusto_bot']:.2f}", "m", "Spessore fusto al piede"),
        ("h_fronte", f"{data['h_fronte']:.2f}", "m", "Altezza terreno a fronte"),
    ]
    for row_data in geo_rows:
        row = t.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)

    _aggiungi_heading2(doc, "1.3.2 Materiali e carichi")
    t2 = doc.add_table(rows=1, cols=4)
    t2.style = 'Table Grid'
    hdr2 = t2.rows[0].cells
    hdr2[0].text = "Parametro"
    hdr2[1].text = "Valore"
    hdr2[2].text = "Unità"
    hdr2[3].text = "Descrizione"
    _intestazione_tabella(t2)

    mat_rows = [
        ("γ_cls", f"{data['gamma_cls']:.1f}", "kN/m³", "Peso di volume calcestruzzo"),
        ("q", f"{data['q']:.1f}", "kPa", "Sovraccarico distribuito a tergo"),
        ("μ", f"{data['mu_base']:.3f}", "–", "Coefficiente di attrito alla base"),
        ("δ", f"{data['delta_muro']:.1f}", "°", "Angolo di attrito terra-muro"),
        ("q_amm", f"{data['q_amm']:.1f}", "kPa", "Pressione ammissibile di riferimento"),
    ]
    for row_data in mat_rows:
        row = t2.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)

    _aggiungi_heading2(doc, "1.3.3 Parametri sismici")
    t3 = doc.add_table(rows=1, cols=4)
    t3.style = 'Table Grid'
    hdr3 = t3.rows[0].cells
    hdr3[0].text = "Parametro"
    hdr3[1].text = "Valore"
    hdr3[2].text = "Unità"
    hdr3[3].text = "Descrizione"
    _intestazione_tabella(t3)

    sis_rows = [
        ("k_h", f"{data['kh']:.3f}", "–", "Coefficiente sismico orizzontale (NTC2018 §7.11.5)"),
        ("k_v", f"{data['kv']:.3f}", "–", "Coefficiente sismico verticale (NTC2018 §7.11.5)"),
    ]
    for row_data in sis_rows:
        row = t3.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)

    if data.get('ha_tirante'):
        _aggiungi_heading2(doc, "1.3.4 Tirante di ancoraggio")
        t4 = doc.add_table(rows=1, cols=4)
        t4.style = 'Table Grid'
        hdr4 = t4.rows[0].cells
        hdr4[0].text = "Parametro"
        hdr4[1].text = "Valore"
        hdr4[2].text = "Unità"
        hdr4[3].text = "Descrizione"
        _intestazione_tabella(t4)
        tir_rows = [
            ("T", f"{data['t_tiro']:.1f}", "kN/m", "Tiro di progetto per metro lineare"),
            ("α", f"{data['t_inclinazione']:.1f}", "°", "Inclinazione rispetto all'orizzontale"),
            ("y_t", f"{data['t_quota']:.2f}", "m", "Quota di applicazione dal fondo scavo"),
        ]
        for row_data in tir_rows:
            row = t4.add_row().cells
            for i, val in enumerate(row_data):
                row[i].text = str(val)

    _aggiungi_heading2(doc, "1.3.5 Stratigrafia adottata")
    strat_df = r.get("stratigrafia")
    if strat_df is not None and not strat_df.empty:
        t_strat = doc.add_table(rows=1, cols=6)
        t_strat.style = 'Table Grid'
        for i, header in enumerate(["Strato", "Spessore [m]", "gamma dry", "gamma sat", "phi [deg]", "cu [kPa]"]):
            t_strat.rows[0].cells[i].text = header
        _intestazione_tabella(t_strat)
        for idx, row_data in strat_df.iterrows():
            row = t_strat.add_row().cells
            row[0].text = str(idx + 1)
            row[1].text = f"{row_data['spessore_m']:.2f}"
            row[2].text = f"{row_data['gamma_dry']:.2f}"
            row[3].text = f"{row_data['gamma_sat']:.2f}"
            row[4].text = f"{row_data['phi_deg']:.1f}"
            row[5].text = f"{row_data['cu_kPa']:.1f}"
    else:
        doc.add_paragraph("Stratigrafia non disponibile nel risultato di calcolo.")

    _aggiungi_heading2(doc, "1.3.6 Elaborati grafici")
    try:
        from src import figura_geometria, figura_output, figura_pendio

        dati_muro = _dati_muro()
        _aggiungi_immagine(
            doc,
            "Figura 1 - Schema geometrico, terreno e falda",
            _plot_png_bytes(figura_geometria(dati_muro)),
        )
        _aggiungi_immagine(
            doc,
            "Figura 2 - Diagrammi delle pressioni attive e passive",
            _plot_png_bytes(figura_output(r)),
        )
        if data.get("pendio"):
            _aggiungi_immagine(
                doc,
                "Figura 3 - Superficie critica per stabilita globale",
                _plot_png_bytes(figura_pendio(dati_muro, data["pendio"])),
            )
    except Exception as exc:
        doc.add_paragraph(
            "Immagini non disponibili in questa esecuzione. Installare `kaleido` per esportare "
            f"automaticamente i plot Plotly nel report. Dettaglio tecnico: {exc}"
        )

    doc.add_paragraph()

    _aggiungi_heading2(doc, "1.4 Tabelle di calcolo")

    _aggiungi_heading2(doc, "1.4.1 Spinte sul muro")
    t_sp = doc.add_table(rows=1, cols=4)
    t_sp.style = 'Table Grid'
    hdr_sp = t_sp.rows[0].cells
    hdr_sp[0].text = "Parametro"
    hdr_sp[1].text = "Statico (GEO)"
    hdr_sp[2].text = "Sismico kv+"
    hdr_sp[3].text = "Unità"
    _intestazione_tabella(t_sp)
    sp_rows = [
        ("Pa — Spinta attiva orizzontale", f"{r['statico']['Pa']:.2f}", f"{r['sismico']['Pa']:.2f}", "kN/m"),
        ("Pp — Spinta passiva", f"{r['statico']['Pp']:.2f}", f"{r['sismico']['Pp']:.2f}", "kN/m"),
    ]
    for row_data in sp_rows:
        row = t_sp.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)

    _aggiungi_heading2(doc, "1.4.2 Verifica a ribaltamento (EQU) - NTC2018 §6.5.3.1.1")
    doc.add_paragraph(
        "Il fattore di sicurezza a ribaltamento è calcolato rispetto al polo posizionato al lembo "
        "estremo di valle (piede della fondazione). Combinazione EQU: γ_dest=1.1, γ_stab=0.9."
    )
    t_rib = doc.add_table(rows=1, cols=4)
    t_rib.style = 'Table Grid'
    hdr_rib = t_rib.rows[0].cells
    hdr_rib[0].text = "Parametro"
    hdr_rib[1].text = "Valore"
    hdr_rib[2].text = "Unità"
    hdr_rib[3].text = "Descrizione"
    _intestazione_tabella(t_rib)

    rib_rows = [
        ("FS_rib (EQU — Statico)", f"{r['st_EQU']['FS_rib']:.3f}", "–", "FS ribaltamento statico (limite ≥ 1.0)"),
        ("FS_rib (Sismico kv+)", f"{r['sismico']['FS_rib']:.3f}", "–", "FS ribaltamento sismico (limite ≥ 1.0)"),
    ]
    for row_data in rib_rows:
        row = t_rib.add_row().cells
        row[0].text = row_data[0]
        _esito_cell(row[1], float(row_data[1]), 1.0)
        row[2].text = row_data[2]
        row[3].text = row_data[3]

    _aggiungi_heading2(doc, "1.4.3 Verifica a scorrimento (GEO) - NTC2018 §6.5.3.1.2")
    doc.add_paragraph(
        "La resistenza allo scorrimento è calcolata per criterio di Mohr-Coulomb all'interfaccia "
        "fondazione-terreno: R_d = μ·V_d + c_u·B + P_p. Combinazione GEO."
    )
    t_sco = doc.add_table(rows=1, cols=4)
    t_sco.style = 'Table Grid'
    hdr_sco = t_sco.rows[0].cells
    hdr_sco[0].text = "Parametro"
    hdr_sco[1].text = "Valore"
    hdr_sco[2].text = "Unità"
    hdr_sco[3].text = "Descrizione"
    _intestazione_tabella(t_sco)
    sco_rows = [
        ("FS_scorr (GEO — Statico)", f"{r['statico']['FS_scorr']:.3f}", "–", "FS scorrimento statico (limite ≥ 1.0)"),
        ("FS_scorr (Sismico kv+)", f"{r['sismico']['FS_scorr']:.3f}", "–", "FS scorrimento sismico (limite ≥ 1.0)"),
    ]
    for row_data in sco_rows:
        row = t_sco.add_row().cells
        row[0].text = row_data[0]
        _esito_cell(row[1], float(row_data[1]), 1.0)
        row[2].text = row_data[2]
        row[3].text = row_data[3]

    _aggiungi_heading2(doc, "1.4.4 Pressioni di contatto e capacita portante (GEO) - NTC2018 §6.4.2")
    doc.add_paragraph(
        "La distribuzione delle pressioni sul piano di posa è calcolata con la formula di Navier. "
        "La capacità portante limite è determinata con la formula trinomia di Brinch-Hansen "
        "(EC7-1 Allegato D / NTC2018 §6.4.2.2)."
    )
    t_por = doc.add_table(rows=1, cols=4)
    t_por.style = 'Table Grid'
    hdr_por = t_por.rows[0].cells
    hdr_por[0].text = "Parametro"
    hdr_por[1].text = "Valore"
    hdr_por[2].text = "Unità"
    hdr_por[3].text = "Descrizione"
    _intestazione_tabella(t_por)
    por_rows = [
        ("q_max (Statico)", f"{r['statico']['qmax']:.2f}", "kPa", "Pressione massima (formula Navier)"),
        ("q_min (Statico)", f"{r['statico']['qmin']:.2f}", "kPa", "Pressione minima (nucleo centrale)"),
        ("q_lim Hansen (Statico)", f"{r['statico']['q_lim']:.2f}", "kPa", "Carico limite Brinch-Hansen"),
        ("FS_portanza (Statico)", f"{r['statico']['FS_portanza']:.3f}", "–", "FS portanza (limite ≥ 1.0)"),
        ("FS_portanza (Sismico)", f"{r['sismico']['FS_portanza']:.3f}", "–", "FS portanza sismico"),
    ]
    for i_r, row_data in enumerate(por_rows):
        row = t_por.add_row().cells
        row[0].text = row_data[0]
        if i_r in (3, 4):
            _esito_cell(row[1], float(row_data[1]), 1.0)
        else:
            row[1].text = row_data[1]
        row[2].text = row_data[2]
        row[3].text = row_data[3]

    doc.add_paragraph()

    _aggiungi_heading2(doc, "1.4.5 Sollecitazioni strutturali allo spiccato")

    doc.add_paragraph(
        "Le sollecitazioni sono calcolate alla quota dello spiccato della fondazione "
        "(y = t_base), per il dimensionamento delle armature alla base del fusto."
    )

    t_sol = doc.add_table(rows=1, cols=5)
    t_sol.style = 'Table Grid'
    hdr_sol = t_sol.rows[0].cells
    hdr_sol[0].text = "Combinazione"
    hdr_sol[1].text = "N [kN/m]"
    hdr_sol[2].text = "V [kN/m]"
    hdr_sol[3].text = "M [kNm/m]"
    hdr_sol[4].text = "Descrizione"
    _intestazione_tabella(t_sol)

    comb_map = [
        ('st_EQU', "Statica EQU"),
        ('statico', "Statica GEO"),
        ('sismico', "Sismica kv+"),
        ('se_neg', "Sismica kv−"),
    ]
    for key, nome in comb_map:
        rd = r[key]
        row = t_sol.add_row().cells
        row[0].text = nome
        row[1].text = f"{rd['N_spiccato']:.2f}"
        row[2].text = f"{rd['V_spiccato']:.2f}"
        row[3].text = f"{rd['M_spiccato']:.2f}"
        row[4].text = "Valori per dim. armature mensola"

    doc.add_paragraph()

    _aggiungi_heading2(doc, "1.4.6 Riepilogo verifiche")

    t_riepilogo = doc.add_table(rows=1, cols=4)
    t_riepilogo.style = 'Table Grid'
    hdr_r = t_riepilogo.rows[0].cells
    hdr_r[0].text = "Verifica"
    hdr_r[1].text = "FS / D.C."
    hdr_r[2].text = "Limite"
    hdr_r[3].text = "Esito"
    _intestazione_tabella(t_riepilogo)

    verifiche = [
        ("Ribaltamento (EQU — Statico)", r['st_EQU']['FS_rib'], "≥ 1.0"),
        ("Ribaltamento (Sismico kv+)", r['sismico']['FS_rib'], "≥ 1.0"),
        ("Scorrimento (GEO — Statico)", r['statico']['FS_scorr'], "≥ 1.0"),
        ("Scorrimento (Sismico kv+)", r['sismico']['FS_scorr'], "≥ 1.0"),
        ("Portanza (GEO — Statico)", r['statico']['FS_portanza'], "≥ 1.0"),
        ("Portanza (Sismico kv+)", r['sismico']['FS_portanza'], "≥ 1.0"),
    ]

    for nome_v, fs_v, limite_v in verifiche:
        row = t_riepilogo.add_row().cells
        row[0].text = nome_v
        row[1].text = f"{fs_v:.3f}"
        row[2].text = limite_v
        esito_run = row[3].paragraphs[0].add_run()
        if fs_v >= 1.0:
            esito_run.text = "VERIFICATO"
        else:
            esito_run.text = "NON VERIFICATO"
        esito_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        esito_run.bold = True

    doc.add_paragraph()

    _aggiungi_heading2(doc, "1.5 Stabilita globale e note tecniche")
    pendio = data.get("pendio")
    if pendio and pendio.get("statico"):
        _aggiungi_tabella_kv(
            doc,
            [
                ("Metodo", "Fellenius semplificato con ricerca automatica del cerchio critico"),
                ("FS pendio statico", f"{pendio['statico']['FS']:.3f}"),
                ("FS pendio sismico", f"{pendio['sismico']['FS']:.3f}" if pendio.get("sismico") else "n.d."),
                ("phi equivalente", f"{pendio['props']['phi_deg']:.1f} deg"),
                ("cu equivalente", f"{pendio['props']['cu_kPa']:.1f} kPa"),
            ],
        )
    else:
        doc.add_paragraph("La verifica globale del pendio non e stata eseguita o non ha restituito una superficie critica valida.")

    if data.get("notes"):
        _aggiungi_heading2(doc, "Note automatiche")
        for note in data["notes"]:
            doc.add_paragraph(str(note), style="List Bullet")

    _aggiungi_heading2(doc, "1.6 Disclaimer")
    p_disc = doc.add_paragraph(
        "I risultati prodotti dall'applicazione CivilBox — Muro di Sostegno sono elaborati "
        "automaticamente sulla base dei dati inseriti dall'utente. L'applicazione non sostituisce "
        "il giudizio professionale di un ingegnere abilitato. I risultati devono essere verificati "
        "criticamente prima di essere utilizzati per qualsiasi scopo progettuale, costruttivo o "
        "autorizzativo. Anthropic e gli autori dell'applicazione declinano ogni responsabilità "
        "per un uso improprio dei risultati."
    )
    p_disc.runs[0].font.size = Pt(9)

    # Salva in BytesIO
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# PDF REPORT
# ---------------------------------------------------------------------------

def create_pdf_report(data: Dict[str, Any]) -> bytes:
    """
    Genera la relazione tecnica in formato PDF (.pdf).
    Restituisce i bytes del file, senza scrivere su disco.
    """
    from fpdf import FPDF

    class PDF(FPDF):
        def header(self):
            self.set_fill_color(44, 62, 80)
            self.rect(0, 0, 210, 12, 'F')
            self.set_font('Helvetica', 'B', 10)
            self.set_text_color(255, 255, 255)
            self.cell(0, 10, 'RELAZIONE TECNICA - MURO DI SOSTEGNO  |  NTC2018 / EC7', align='C')
            self.ln(14)
            self.set_text_color(0, 0, 0)

        def footer(self):
            self.set_y(-12)
            self.set_font('Helvetica', 'I', 8)
            self.set_text_color(120, 120, 120)
            self.cell(0, 10, f'Pagina {self.page_no()} - CivilBox - Muro di Sostegno - {date.today()}', align='C')

    def _sezione(pdf, testo):
        pdf.set_text_color(0, 0, 0)
        pdf.set_font('Helvetica', 'B', 13)
        pdf.cell(0, 8, testo, ln=True)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.set_text_color(0, 0, 0)
        pdf.ln(2)

    def _sottosezione(pdf, testo):
        pdf.set_font('Helvetica', 'B', 10)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(0, 6, testo, ln=True)
        pdf.set_text_color(0, 0, 0)
        pdf.ln(1)

    def _tabella(pdf, intestazioni, righe, larghezze):
        pdf.set_fill_color(229, 231, 235)
        pdf.set_text_color(0, 0, 0)
        pdf.set_font('Helvetica', 'B', 9)
        for i, (col, w) in enumerate(zip(intestazioni, larghezze)):
            pdf.cell(w, 7, col, border=1, fill=True)
        pdf.ln()
        pdf.set_text_color(0, 0, 0)
        alt_row = False
        for riga in righe:
            if alt_row:
                pdf.set_fill_color(240, 245, 250)
            else:
                pdf.set_fill_color(255, 255, 255)
            pdf.set_font('Helvetica', '', 8)
            for i, (val, w) in enumerate(zip(riga, larghezze)):
                pdf.cell(w, 6, str(val), border=1, fill=True)
            pdf.ln()
            alt_row = not alt_row
        pdf.ln(2)

    def _plot_png_file(fig, width=1100, height=700):
        import plotly.io as pio

        fig.update_layout(font=dict(color="#000000"), paper_bgcolor="white", plot_bgcolor="white")
        png = pio.to_image(fig, format="png", width=width, height=height, scale=2)
        tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        tmp.write(png)
        tmp.close()
        return tmp.name

    def _immagine_pdf(pdf, titolo, fig):
        _sottosezione(pdf, titolo)
        path = _plot_png_file(fig)
        try:
            pdf.image(path, x=20, w=170)
            pdf.ln(3)
        finally:
            Path(path).unlink(missing_ok=True)

    def _esito_riga(pdf, nome, fs, limite=1.0):
        pdf.set_font('Helvetica', '', 8)
        pdf.set_fill_color(255, 255, 255)
        pdf.cell(90, 6, nome, border=1, fill=True)
        pdf.cell(25, 6, f"{fs:.3f}", border=1, fill=True)
        pdf.cell(20, 6, f">= {limite:.1f}", border=1, fill=True)
        if fs >= limite:
            pdf.set_fill_color(255, 255, 255)
            pdf.set_text_color(0, 0, 0)
            pdf.cell(35, 6, "VERIFICATO", border=1, fill=True, align='C')
        else:
            pdf.set_fill_color(255, 255, 255)
            pdf.set_text_color(0, 0, 0)
            pdf.cell(35, 6, "NON VERIFICATO", border=1, fill=True, align='C')
        pdf.set_text_color(0, 0, 0)
        pdf.ln()

    pdf = PDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    r = data['risultati']
    warnings = data.get("warnings", [])
    q_ratio_statico = r['statico']['qmax'] / max(data['q_amm'], 1e-9)

    _sezione(pdf, "1. RELAZIONE DI CALCOLO - MURO DI SOSTEGNO")
    _sottosezione(pdf, "1.1 Frontespizio")
    pdf.set_font('Helvetica', '', 9)
    pdf.cell(50, 6, "Applicazione:", ln=False)
    pdf.set_font('Helvetica', 'B', 9)
    pdf.cell(0, 6, "Muro - Analisi Avanzata NTC2018", ln=True)
    pdf.set_font('Helvetica', '', 9)
    pdf.cell(50, 6, "Normative:", ln=False)
    pdf.cell(0, 6, "NTC2018 §6.5.3, EC7-1 §6.5, EC8-5 §7.3", ln=True)
    pdf.cell(50, 6, "Data:", ln=False)
    pdf.cell(0, 6, str(date.today()), ln=True)
    pdf.ln(4)

    _sottosezione(pdf, "1.2 Sintesi esecutiva")
    _tabella(pdf,
        ["Controllo", "Valore", "Esito"],
        [
            ("Ribaltamento statico EQU", f"FS {r['st_EQU']['FS_rib']:.3f}", "VERIFICATO" if r['st_EQU']['FS_rib'] >= 1.0 else "NON VERIFICATO"),
            ("Scorrimento statico GEO", f"FS {r['statico']['FS_scorr']:.3f}", "VERIFICATO" if r['statico']['FS_scorr'] >= 1.0 else "NON VERIFICATO"),
            ("Portanza statica GEO", f"FS {r['statico']['FS_portanza']:.3f}", "VERIFICATO" if r['statico']['FS_portanza'] >= 1.0 else "NON VERIFICATO"),
            ("q_max / q_amm", f"{q_ratio_statico:.3f}", "VERIFICATO" if q_ratio_statico <= 1.0 else "NON VERIFICATO"),
            ("Criticita automatiche", str(len(warnings)), "DA ESAMINARE" if warnings else "NESSUNA"),
        ],
        [75, 45, 50]
    )

    _sottosezione(pdf, "1.3 Dati di input - geometria e carichi")
    _tabella(pdf,
        ["Parametro", "Valore", "Unità", "Descrizione"],
        [
            ("H", f"{data['H']:.2f}", "m", "Altezza terreno a tergo"),
            ("B", f"{data['B']:.2f}", "m", "Base totale fondazione"),
            ("B_punta / B_tallone", f"{data['B_punta']:.2f} / {data['B_tallone']:.2f}", "m", "Punta e tallone"),
            ("t_base / t_fusto", f"{data['t_base']:.2f} / {data['t_fusto_bot']:.2f}", "m", "Spessori base e fusto"),
            ("gamma_cls", f"{data['gamma_cls']:.1f}", "kN/m³", "Peso volume calcestruzzo"),
            ("q", f"{data['q']:.1f}", "kPa", "Sovraccarico a tergo"),
            ("mu_base", f"{data['mu_base']:.3f}", "-", "Coeff. attrito base"),
            ("delta", f"{data['delta_muro']:.1f}", "°", "Attrito terra-muro"),
            ("kh / kv", f"{data['kh']:.3f} / {data['kv']:.3f}", "-", "Coeff. sismici pseudo-statici"),
            ("falda_retro", f"{data['falda_retro']:.1f}", "m", "Profondità falda a tergo"),
        ],
        [35, 30, 20, 85]
    )

    try:
        from src import DatiMuro, DEFAULT_STRAT, figura_geometria, figura_output, figura_pendio

        dati_muro = DatiMuro(
            H=data["H"],
            q=data["q"],
            gamma_cls=data["gamma_cls"],
            B=data["B"],
            B_punta=data["B_punta"],
            B_tallone=data["B_tallone"],
            t_base=data["t_base"],
            t_fusto_top=data["t_fusto_top"],
            t_fusto_bot=data["t_fusto_bot"],
            mu_base=data["mu_base"],
            q_amm=data["q_amm"],
            h_fronte=data["h_fronte"],
            falda_retro=data["falda_retro"],
            falda_fronte=data["falda_fronte"],
            kh=data["kh"],
            kv=data["kv"],
            delta_muro=data["delta_muro"],
            include_passivo=data.get("include_passivo", True),
            stratigrafia_csv=data.get("stratigrafia_csv", DEFAULT_STRAT),
            ha_tirante=data.get("ha_tirante", False),
            t_quota=data.get("t_quota", 0.0),
            t_inclinazione=data.get("t_inclinazione", 0.0),
            t_tiro=data.get("t_tiro", 0.0),
            analizza_pendio=bool(data.get("pendio")),
            pendio_H=data.get("pendio_H", 5.0),
            pendio_beta=data.get("pendio_beta", 26.0),
            pendio_berma=data.get("pendio_berma", 6.0),
        )
        _sottosezione(pdf, "1.4 Elaborati grafici")
        _immagine_pdf(pdf, "Figura 1 - Schema geometrico, terreno e falda", figura_geometria(dati_muro))
        _immagine_pdf(pdf, "Figura 2 - Diagrammi delle pressioni attive e passive", figura_output(r))
        if data.get("pendio"):
            _immagine_pdf(pdf, "Figura 3 - Superficie critica per stabilita globale", figura_pendio(dati_muro, data["pendio"]))
    except Exception as exc:
        pdf.set_font('Helvetica', '', 8)
        pdf.multi_cell(0, 5, f"Immagini non disponibili in questa esecuzione. Installare kaleido. Dettaglio tecnico: {exc}")
        pdf.ln(2)

    _sottosezione(pdf, "1.5 Tabelle di calcolo - verifiche geotecniche")
    _tabella(pdf,
        ["Grandezza", "Statico (GEO)", "Sismico kv+", "Unità"],
        [
            ("Pa - Spinta attiva [kN/m]", f"{r['statico']['Pa']:.2f}", f"{r['sismico']['Pa']:.2f}", "kN/m"),
            ("Pp - Spinta passiva [kN/m]", f"{r['statico']['Pp']:.2f}", f"{r['sismico']['Pp']:.2f}", "kN/m"),
            ("q_max [kPa]", f"{r['statico']['qmax']:.2f}", f"{r['sismico']['qmax']:.2f}", "kPa"),
            ("q_min [kPa]", f"{r['statico']['qmin']:.2f}", f"{r['sismico']['qmin']:.2f}", "kPa"),
            ("q_lim Hansen [kPa]", f"{r['statico']['q_lim']:.2f}", f"{r['sismico']['q_lim']:.2f}", "kPa"),
        ],
        [65, 35, 35, 35]
    )

    _sottosezione(pdf, "1.5.1 Sollecitazioni strutturali allo spiccato")
    _tabella(pdf,
        ["Combinazione", "N [kN/m]", "V [kN/m]", "M [kNm/m]"],
        [
            ("Statica EQU", f"{r['st_EQU']['N_spiccato']:.2f}", f"{r['st_EQU']['V_spiccato']:.2f}", f"{r['st_EQU']['M_spiccato']:.2f}"),
            ("Statica GEO", f"{r['statico']['N_spiccato']:.2f}", f"{r['statico']['V_spiccato']:.2f}", f"{r['statico']['M_spiccato']:.2f}"),
            ("Sismica kv+", f"{r['sismico']['N_spiccato']:.2f}", f"{r['sismico']['V_spiccato']:.2f}", f"{r['sismico']['M_spiccato']:.2f}"),
            ("Sismica kv-", f"{r['se_neg']['N_spiccato']:.2f}", f"{r['se_neg']['V_spiccato']:.2f}", f"{r['se_neg']['M_spiccato']:.2f}"),
        ],
        [60, 40, 40, 30]
    )

    _sottosezione(pdf, "1.5.2 Riepilogo verifiche")
    pdf.set_font('Helvetica', 'B', 9)
    pdf.set_fill_color(229, 231, 235)
    pdf.set_text_color(0, 0, 0)
    for hdr_txt, w in [("Verifica", 90), ("FS", 25), ("Limite", 20), ("Esito", 35)]:
        pdf.cell(w, 7, hdr_txt, border=1, fill=True)
    pdf.ln()

    verifiche_pdf = [
        ("Ribaltamento (EQU - Statico)", r['st_EQU']['FS_rib']),
        ("Ribaltamento (Sismico kv+)", r['sismico']['FS_rib']),
        ("Scorrimento (GEO - Statico)", r['statico']['FS_scorr']),
        ("Scorrimento (Sismico kv+)", r['sismico']['FS_scorr']),
        ("Portanza (GEO - Statico)", r['statico']['FS_portanza']),
        ("Portanza (Sismico kv+)", r['sismico']['FS_portanza']),
    ]
    for nome_v, fs_v in verifiche_pdf:
        _esito_riga(pdf, nome_v, fs_v)
    pdf.ln(4)

    pendio = data.get("pendio")
    _sottosezione(pdf, "1.6 Stabilita globale")
    if pendio and pendio.get("statico"):
        _tabella(pdf,
            ["Parametro", "Valore", "Unita"],
            [
                ("FS pendio statico", f"{pendio['statico']['FS']:.3f}", "-"),
                ("FS pendio sismico", f"{pendio['sismico']['FS']:.3f}" if pendio.get("sismico") else "n.d.", "-"),
                ("phi equivalente", f"{pendio['props']['phi_deg']:.1f}", "deg"),
                ("cu equivalente", f"{pendio['props']['cu_kPa']:.1f}", "kPa"),
            ],
            [70, 45, 30]
        )
    else:
        pdf.set_font('Helvetica', '', 8)
        pdf.multi_cell(0, 5, "Verifica globale del pendio non eseguita o senza superficie critica valida.")
        pdf.ln(2)

    # Disclaimer
    _sottosezione(pdf, "1.7 Disclaimer")
    pdf.set_font('Helvetica', 'I', 8)
    pdf.multi_cell(
        0, 5,
        "I risultati prodotti dall'applicazione CivilBox - Muro di Sostegno sono elaborati "
        "automaticamente sulla base dei dati inseriti dall'utente. L'applicazione non sostituisce "
        "il giudizio professionale di un ingegnere abilitato. I risultati devono essere verificati "
        "criticamente prima di essere utilizzati per qualsiasi scopo progettuale, costruttivo o "
        "autorizzativo."
    )

    pdf_bytes = pdf.output()
    if isinstance(pdf_bytes, str):
        return pdf_bytes.encode('latin-1', errors='ignore')
    return bytes(pdf_bytes)
